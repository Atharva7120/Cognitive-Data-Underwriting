import re, time
from unidecode import unidecode
import docx
import sys
import nltk
import codecs
from nltk.tokenize import PunktSentenceTokenizer
from nltk.corpus import stopwords
from docx.shared import Inches
import os
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import shutil
from pymongo import MongoClient
import psycopg2
from nltk.util import ngrams
from nltk.corpus.reader import WordListCorpusReader
from nltk.tokenize import sent_tokenize
import datetime



def getText(file_path):
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\n'.join(text)

def extract_fields(text):
    policy = "not captured"
    dt1 = "not captured"
    dob1 = "not captured"
    dob = None
    dt = None
    yr = None
    ft = 999

    for line in text.split('\n'):
        string7 = re.search(r"(#|Date of Birth:|Date of Birth)\s*:\s*(\d{2}/\d{2}/\d+)", line)
        if string7:
            dob = string7.group(2)
            dob1 = dob

        string8 = re.search(r"(Client:|arc edc som)\s*:\s*(\d{2}[-/]\d{2}[-/]\d{2})", line)
        if string8:
            dt = string8.group(2)
            dt1 = dt
            if dt:
                yr = dt.split("/")[-1]
                ft = yr

        string9 = re.search(r"Policy#:\s*(\d+-\d+)", line)
        if string9:
            policy = string9.group(1)

    return policy, dt1, dob1, dob, dt, yr, ft


def remove_sentence_vit(text, word):
    sentences = sent_tokenize(text)
    for sentence in sentences:
        if word in list_words_vitals:
            if (all((word in sentence, word))):
                string = re.search(r'height|ht|height:|ht:', sentence, re.M | re.I)
                if string is not None:
                    height = sentence
                    string3 = re.findall(r'\d+[\.\d{1,2}]?\s*in |\d*\s*ft\s*\d*\s*in|\d*\s*in', height,
                                         re.M | re.I)
                    if string3 is not None:
                        for k in string3:
                            f5.write("height: " + k + "\n")

                string1 = re.search(r'weight|wt|weight:|wt:', sentence, re.M | re.I)
                if string1 is not None:
                    weight = sentence
                    string4 = re.findall(r'\d+[\.\d{1,2}]*\s*lbs|\d+[\.\d{1,2}]*\s*lb|\d*\s*pounds', weight,
                                         re.M | re.I)

                    if string4 is not None:
                        for k in string4:
                            f5.write("weight: " + k + "\n")

                string2 = re.search(r'blood pressure|bp|blood pressure:|bp:', sentence, re.M | re.I)
                if string2 is not None:
                    bloodp = sentence
                    string5 = re.findall(r'\s*\d+\s*/\s*\d+\s*', bloodp, re.M | re.I)
                    if string5 is not None:
                        for k in string5:
                            f5.write("blood pressure: " + k + "\n")



def remove_sentence_date(text, word):
    sentences = sent_tokenize(text)
    for sentence in sentences:
        if word in list_dates:
            if (all((word in sentence, word))):
                f5.write(sentence + "\n")


def remove_sentence_sum(text, word):
        sentences = sent_tokenize(text)
        for sentence in sentences:
            if word in list_words + list_words_dis:
                if (all((word in sentence, word))):
                    f5.write(sentence + "\n")


def remove_sentence_so(text, word):
    sentences = sent_tokenize(text)
    for sentence in sentences:
        if word in list_words_So_history:
            if (all((word in sentence, word))):
                f5.write(sentence + "\n")


def remove_sentence_fam(text, word):
    sentences = sent_tokenize(text)
    for sentence in sentences:
        if word in list_words_Fa_history:
            if (all((word in sentence, word))):
                f5.write(sentence + "\n")



ALLOWED_FORMATS = ['%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d', '%d.%m.%Y', '%d-%m-%y', '%m.%d.%Y', '%m/%d/%Y',
                   '%m-%d-%y', '%b/%d/%Y', '%b %d %Y', '%m/%d/%y', '%B %d %Y']


def convert_date(string):
    for format in ALLOWED_FORMATS:
        try:
            conv = time.strptime(string, format)
            return time.strftime('%m/%d/%y', conv)
        except ValueError:
            pass


def getKey(item):
    dates.append(item)
    for i in range(len(dates[0])):
        for j in range(len(dates[0])):
            if dates[0][i][2] < dates[0][j][2]:
                temp = dates[0][i]
                dates[0][i] = dates[0][j]
                dates[0][j] = temp
            elif dates[0][i][2] == dates[0][j][2]:
                if dates[0][i][0] < dates[0][j][0]:
                    temp1 = dates[0][i]
                    dates[0][i] = dates[0][j]
                    dates[0][j] = temp1
            elif dates[0][i][2] == dates[0][j][2] and dates[0][i][0] == dates[0][j][0]:
                if dates[0][i][1] < dates[0][j][1]:
                    temp2 = dates[0][i]
                    dates[0][i] = dates[0][j]
                    dates[0][j] = temp2


def findWhole(w):
    return re.compile(r'\b({0})\b'.format(w), flags=re.IGNORECASE).search






directory = "E:\BE"
for filename in os.listdir(directory):
    file_path = os.path.join(directory, filename)
    folder=filename.split(".")[0]
    if not (os.path.exists(folder)) and not (os.path.isdir(folder)):
        os.makedirs(folder)
    folder_path=os.path.join(folder, filename)
    if filename.endswith('.docx'):
        if os.path.exists(folder_path):
            try:
                os.remove(folder_path)
                print(f"The file '{folder_path}' has been successfully removed.")
            except Exception as e:
                print(f"An error occurred while trying to remove the file: {e}")
        else:
            text = getText(file_path)
            with open(folder+"/"+filename[:-5] + '.txt', 'w', encoding='utf-8') as text_file:
                text_file.write(text)
            print(f'Converted {filename} to text')

            text = text.replace('$', 's')

            # Save the modified text to a new text file
            with open(folder + "/" + filename[:-5] + '.txt', 'w', encoding='utf-8') as text_file:
                text_file.write(text)

            print(f'Converted {filename} to text and replaced "$" with "s".')

            # Extract fields from the modified text
            policy, dt1, dob1, dob, dt, yr, ft = extract_fields(text)

            # You can print or process the extracted fields here
            # print(f'Policy: {policy}')
            # print(f'Date: {dt1}')
            # print(f'Date of Birth: {dob1}')
            # print(f'dob: {dob}')
            # print(f'dt: {dt}')
            # print(f'yr: {yr}')
            # print(f'ft: {ft}')

            str3 = folder + "/" + filename[:-5] + '.txt'
            str4 = folder + "/p1.txt"
            f1 = codecs.open(str4, "w")
            with open(str3) as f:
                for line in f:
                    if re.search("|".join(
                            ["(V[^.]*t)\s*(D[^.]*)\s*(T[^.]*e)", "(Encounter Date: )", "ORDER DATE", "DIS DATE",
                             "Print da[^.]*me:", "Date\s*\d+\/\d+\/\d{4}", "\w{4} \d{2}[,|>| |0-9] \d{3}",
                             "Marc Lee MD", "Date:\s*\d+\/\d+[{0-9}\/]\d+", "^[\d+\/\d+\/\d+][^.]*(AM|PM)",
                             "Last\s*Saved:", "DATE OF SERVICE:"]), line):
                        us = re.findall("|".join(
                            ['\d+\/\d+\/\d+', '\w+\s+\d+\s+\d{4}', '\w+ \d+[,|>| |0-9] \d+', '\d+\/\d+[{0-9}\/]\d+']),
                            line)
                        for i in range(len(us)):
                            de1 = us[i]
                            sp = re.sub("|".join([":", ",", ";", ":: ,", ">", ]), " ", de1)
                            yes = convert_date(sp)
                            if yes != None and yes != dob and yes != dt:
                                if yr != None:
                                    yrr = re.split("/", yes)
                                    if (yrr[0] != yr[0] or yrr[2] != yr[2]):
                                        su = re.sub(
                                            "|".join(['\d+\/\d+\/\d+', '\w+\s+\d+\s+\d+', '\w+ \d+[,|>| |0-9] \d+']),
                                            yes, line)
                                        if len(su) < 100:
                                            tr = "Visit Date/Time " + su
                                            result = re.search('%s(.*)%s' % ("Visit Date/Time", "\d{2}\/\d{2}\/\d{2}"),
                                                               tr).group(1)
                                            sup = tr.replace(result, " ")
                                            f1.write(sup)

                                else:
                                    su = re.sub(
                                        "|".join(['\d+\/\d+\/\d+', '\w+\s+\d+\s+\d+', '\w+ \d+[,|>| |0-9] \d+']), yes,
                                        line)
                                    if len(su) < 100:
                                        tr = "Visit Date/Time " + su
                                        result = re.search('%s(.*)%s' % ("Visit Date/Time", "\d{2}\/\d{2}\/\d{2}"),
                                                           tr).group(1)
                                        sup = tr.replace(result, " ")
                                        f1.write(sup)

                    else:
                        f1.write(line)
            f1.close()
            f = codecs.open(str4, "r")
            str5 = folder + "/punc.txt"
            f1 = codecs.open(str5, "w")
            f2 = f.read()
            lines = []
            for word in f2.splitlines():
                if word != '':
                    lines.append(word)

            for k in range(len(lines)):
                lines[k] = re.sub("\s\s+", " ", lines[k])
                lines[k] = re.sub("\t", " ", lines[k])

            for i in range(len(lines)):
                l = len(lines[i]) - 1
                if lines[i][l] != '.':
                    if (i + 1) < (len(lines) - 1) and (lines[i + 1][0].isupper() or lines[i + 1][0].isnumeric()):
                        f1.write(lines[i].lower())
                        f1.write('. \n')
                    else:
                        f1.write(lines[i].lower() + ' ')
                else:
                    f1.write(lines[i].lower() + ' ')

            f.close()
            f1.close()

            f2 = codecs.open(str5, "r")
            text = f2.read()
            str6 = folder + "/para.txt"
            f3 = codecs.open(str6, "w")
            sp = re.split("visit date/time", text)
            for i in range(len(sp)):
                f3.write("visit date/time ")
                f3.write(sp[i])
                f3.write("\n----------------------------------------------------------------\n")

            f2.close()
            f3.close()


            f7 = codecs.open(str6, "r")
            str7 = folder + "/summary.txt"
            f5 = codecs.open(str7, "w")
            dates = WordListCorpusReader('E:\BE\RetrievalCorpora', ['date.txt'])
            conditions = WordListCorpusReader('E:\BE\RetrievalCorpora',
                                              ['indicative_words.txt'])  # produce list of words from file
            diseases = WordListCorpusReader('E:\BE\RetrievalCorpora', ['diseases.txt'])  # produce list of words from fil
            social = WordListCorpusReader('E:\BE\RetrievalCorpora', ['social.txt'])  # produce list of words from file
            family = WordListCorpusReader('E:\BE\RetrievalCorpora', ['family.txt'])
            vitals = WordListCorpusReader('E:\BE\RetrievalCorpora', ['vitals.txt'])  # produce list of words from file
            list_words_vitals = vitals.words()
            list_dates = dates.words()
            list_words = conditions.words()
            list_words_dis = diseases.words()
            list_words_So_history = social.words()
            list_words_Fa_history = family.words()

            f4 = codecs.open(str6, "r")
            text1 = f4.read()
            sq = re.split("\n----------------------------------------------------------------\n", text1)
            for i in range(len(sq)):
                for word0 in list_dates:
                    remove_sentence_date(sq[i], word0)
                f5.write("\nA] Vitals:\n")
                for word0 in list_words_vitals:
                    remove_sentence_vit(sq[i], word0)
                f5.write("\n")
                f5.write("\nB] Family History:\n")
                for word2 in list_words_Fa_history:
                    remove_sentence_fam(sq[i], word2)
                f5.write("\n")
                f5.write("\nC] Social History:\n")
                for word1 in list_words_So_history:
                    remove_sentence_so(sq[i], word1)
                f5.write("\n")

                f5.write("\nD] Summary:\n")
                for word in list_words + list_words_dis:
                    remove_sentence_sum(sq[i], word)
                f5.write("\n")
            f5.close()
            f7.close()

            f4.close()

            f20 = codecs.open(str7, "r")
            text = f20.read()
            str8 = folder + "/summary1.txt"
            f30 = codecs.open(str8, "w")
            sp = re.split("visit date/time", text)
            for i in range(len(sp)):
                f30.write("visit date/time")
                f30.write(sp[i])
                f30.write("\n----------------------------------------------------------------\n")
            f30.close()
            f20.close()

            input1 = open(str8, 'r')
            str9 = folder + "/final_summary.txt"
            out = open(str9, 'w')
            text1 = input1.read()
            text = re.split('\n----------------------------------------------------------------\n', text1)
            for i in range(len(text)):
                s = set()
                for line in text[i].splitlines():
                    if line not in s:
                        out.write(line)
                        out.write(" \n")
                        s.add(line)

            out.close()
            input1.close()

            f201 = codecs.open(str9, "r")
            text = f201.read()
            str10 = folder + "/final_summary1.txt"
            f301 = codecs.open(str10, "w")
            sp = re.split("visit date/time", text)
            for i in range(len(sp)):
                f301.write("visit date/time")
                f301.write(sp[i])
                f301.write("\n----------------------------------------------------------------\n")
            f301.close()
            f201.close()

            b = []
            v = []
            c = []
            ty = []
            f3 = codecs.open(str10, 'r')
            str11 = folder + "/sort_summary.txt"
            f4 = codecs.open(str11, "w")
            summary = f3.read()
            us1 = re.findall('visit date/time  \d{2}\/\d{2}\/\d{2}', summary)
            reset = set(us1)
            for l in reset:
                v.append(l)
            for i in range(len(v)):
                yes1 = re.findall('\d{2}\/\d{2}\/\d{2}', v[i])
                for j in range(len(yes1)):
                    yes2 = re.split("/", yes1[j])
                    b.append(yes2)
            if ft != 999:
                for i in range(len(b)):
                    if b[i][2] > ft:
                        ty.append(b[i])
                    else:
                        continue
                for l in ty:
                    b.remove(l)

            print(b)
            dates = []

            getKey(b)
            for k in dates[0]:
                t = "/".join(k)
                print(t)
                for i in range(len(v)):
                    if re.search(t, v[i]):
                        gre = summary.split('\n----------------------------------------------------------------\n')
                        for q in gre:
                            if re.search(v[i], q):
                                f4.write(q + "\n\n...........................\n\n")

            f114 = codecs.open(str11, "r")
            str12 = folder + "/formatted_sorted_summary.txt"
            f115 = codecs.open(str12, "w")

            fr = f114.read()
            fr = re.sub(r'([A|B|C|D]])', r'\n\n\1', fr)

            f115.write((fr))

            f114.close()
            f115.close()

            f118 = open(str12, "r")
            str13 = folder + "/output1.txt"
            f119 = open(str13, "w")

            fr5 = f118.read()
            fr5 = re.sub(
                r'visit date/time  \d{2}\/\d{2}\/\d{2}(.*)\s+A] Vitals:\s+B] Family History:\s+C] Social History:\s+D] Summary:\s+(\.)+',
                "", fr5)
            f119.write(fr5)

            f118.close()
            f119.close()

            f116 = open(str13, "r")
            str14 = folder + "/output.txt"
            f117 = open(str14, "w")

            fr4 = f116.read()
            fr4 = re.sub(r'\n\n\n*', r'\n\n', fr4)
            f117.write(fr4)

            f116.close()
            f117.close()

            f45 = open(str14, "r")
            str15 = folder + "/out1.txt"
            f46 = open(str15, "w")
            unique = []
            um = []
            rd = f45.read()
            yl = rd.split("...........................")
            [unique.append(item) for item in yl if item not in unique]
            for i in unique:
                f46.write(i)
                f46.write("...........................")
            f46.close()
            f45.close()

            mylist = []
            str16 = folder + "/final_output.txt"
            f48 = open(str16, "w")
            f48.write("Policy Number (Account Number): " + str(policy) + "\n")
            f48.write("Client Name: XXXXX\n")
            f48.write("Date of Birth: " + str(dob1) + "\n")
            f48.write("Date Received: " + str(dt1) + "\n")
            f48.write("Date Processed: ")
            today = datetime.date.today()
            mylist.append(today)
            for entry in mylist:
                f48.write(str(entry))

            f48.write('\n\n\n')
            f48.write(
                '------------------------------------------------------------------------------------------------------------------\n')
            with open(str15) as fr:
                for line in fr:
                    f48.write(line)
            f48.close()

            path = str16
            document = docx.Document()
            str1 = "Summary Worksheet"

            paragraph = document.add_heading(str1)
            run = paragraph.add_run()
            font = run.font

            font.size = Pt(14)

            run.underline = True

            paragraph.add_run()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.underline = True
            myfile = open(path).read()
            myfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', myfile)  # remove all non-XML-compatible characters
            p = document.add_paragraph(myfile)
            str17 = folder + "/output.docx"
            document.save(str17)

            str17 = folder + "/output.txt"
            f49 = open(str17, "r")
            b = f49.read()
            b1 = []
            c1 = []
            result = [x.split('D] Summary:') for x in b.split('...........................')]
            for i in range(len(result) - 1):
                b1.append(result[i][1])
                c1.append(result[i][0])
            conditions = WordListCorpusReader('E:\BE\RetrievalCorpora',
                                              ['diseases.txt'])  # produce list of disease words from file
            list_words = conditions.words()

            conditions1 = WordListCorpusReader('E:\BE\RetrievalCorpora',
                                               ['neg_sent.txt'])  # produce list of negative words from file
            list_words1 = conditions1.words()
            str18 = folder + "/present_diseases.txt"
            present = open(str18, 'w')
            str19 = folder + "/init_vector.txt"
            vector = open(str19, 'w')

            for k in b1:
                b2 = (sent_tokenize(k))
                for sentence in b2:

                    for dis in list_words:
                        if dis in sentence:
                            present.write(dis)
                            present.write(":")
                            present.write(sentence)
                            present.write('\n\n')
                            vector.write(dis)
                            vector.write(':')
                            for neg_word in list_words1:
                                flag = 1
                                if (findWhole(neg_word)(sentence)) != None:
                                    flag = 0
                                    break

                            if flag == 1:
                                vector.write('Yes')
                                vector.write('\n')
                            else:
                                vector.write('No')
                                vector.write('\n')
            f49.close()
            present.close()
            vector.close()

            str20 = folder + "/final_vector.txt"
            vec_final = open(str20, "w")
            str21 = folder + "/init_vector.txt"
            vec = open(str21, "r")
            vec_read = vec.read()
            lines = vec_read.splitlines()
            dis_list = []
            val_list = []
            for i in lines:
                k = i.split(':')
                dis_list.append(k[0].lower())
                val_list.append(k[1])

            dis_set = set(dis_list)
            for j in dis_set:
                count_yes = 0
                count_no = 0
                for l in range(len(dis_list)):
                    if dis_list[l] == j:
                        if val_list[l] == 'Yes':
                            count_yes += 1
                        else:
                            count_no += 1
                vec_final.write(j + ':')
                if count_yes > count_no:
                    vec_final.write('Yes\n')
                else:
                    vec_final.write('No\n')

            vec.close()
            vec_final.close()

            reader = WordListCorpusReader('E:\BE\RetrievalCorpora', ['subsec.txt'])  # produce list of words from file
            list_words = reader.words()  # nltk.tokenize.line_tokenize() on raw file data
            str24 = folder + "/output.txt"
            f54 = open(str24, "r")
            s = f54.read()
            result = re.findall(r"(%s)([\s\S]*?)(?=%s|$)" % ('|'.join(list_words), '|'.join(list_words)), s)

            d = {}
            for res in result:
                d.setdefault(res[0], "")
                d[res[0]] += res[1]
            str25 = folder + "/fam_hist.txt"
            f55 = open(str25, "w")
            for key in d:
                if key == "B] Family History:":
                    f55.write(key + "".join(d[key]) + "========================\n")
            str26 = folder + "/social_hist.txt"
            f56 = open(str26, "w")
            for key in d:
                if key == "C] Social History:":  # all parameters required can be included
                    f56.write(key + "".join(d[key]) + "========================\n")

            f54.close()
            f55.close()
            f56.close()

            diseases = WordListCorpusReader('E:\BE\RetrievalCorpora', ['diseases.txt'])
            list_diseases = diseases.words()

#----------------------------------------------------------------------------------------------------------
            try:
                conn = psycopg2.connect(database="postgres", user="postgres", host="localhost", password="post123",
                                        port="5432")
            except:
                print("I am unable to connect to the database")
            cur = conn.cursor()
            cur.execute("select exists(select * from information_schema.tables where table_name=%s)",
                        ('client_records',))
            if cur.fetchone()[0] == True:
                print("Table exists")
            else:
                cur.execute(
                    "create table client_records(policy_num varchar PRIMARY KEY,dob varchar,gender text,visit_date varchar,height_cm varchar, weight_kg varchar,bp varchar,smoking_social_hist text,frequency_smoking text,alcohol_social_hist text,fam_hist_mother text,fam_hist_father text);")
                for d in list_diseases:
                    cur.execute("""alter table client_records add column "%s" text DEFAULT ('No');""" % d)
                cur.execute("Alter table client_records add column risk_class text")
            conn.commit()
            policy_num = "NA"
            dob = "NA"
            gender = "NA"
            visit = "NA"
            hei = "NA"
            wei = "NA"
            bp = "NA"
            hei1 = "NA"
            wei1 = "NA"
            str22 = folder + "/final_output.txt"
            f52 = open(str22, "r")
            f2 = f52.read()
            f3 = f2.splitlines()
            for i in range(len(f3)):

                string7 = re.search("Date of Birth:", f3[i])
                if string7 != None:
                    s1 = re.findall('\d{2}[-/]\d{2}[-/]\d{2}', f3[i])
                    for k in s1:
                        dob = k
                string9 = re.search("Policy Number ", f3[i])
                if string9 != None:
                    s6 = re.findall("\d+[-]\d+", f3[i])
                    for k in s6:
                        policy_num = k
                string10 = re.search("visit date/time ", f3[i])
                if string10 != None:
                    s7 = re.findall("\d{2}[-/]\d{2}[-/]\d{2}", f3[i])
                    for k in s7:
                        visit = k
                string11 = re.search("height:", f3[i])
                if string11 != None:
                    s2 = re.findall("\d+ in|\d+ ft \d+ in|\d+ In", f3[i])
                    for k in s2:
                        hei = k
                string13 = re.search("weight:", f3[i])
                if string13 != None:
                    s3 = re.findall("\d+\s*\w+|\d*\s*ft\s*\d*\s*in", f3[i])
                    for k in s3:
                        wei = k
                string13 = re.search("blood pressure:", f3[i])
                if string13 != None:
                    s4 = re.findall("\d+[/]\d+", f3[i])
                    for k in s4:
                        bp = k
            if wei != "NA":
                new = re.findall("\d+", wei)
                ori1 = float(new[0])
                wei1 = ori1 * 0.453592
            if hei != "NA":
                new1 = re.findall("\w+", hei)
                if new1[1] == "In" or new1[1] == "in":
                    ori_in = float(new1[0])
                    hei1 = ori_in * 2.54  # in to cm
                if new1[1] == "ft" and new1[3] == "in":
                    ori = float(new1[0])  # number ft
                    new = ori * 30.48  # ft to cm
                    ori_in = float(new1[2])
                    new_in = ori_in * 2.54  # num in to cm
                    hei1 = new + new_in

            f52.close()
            str23 = folder + "/punc.txt"
            f53 = open(str23, "r")
            f21 = f53.read()
            f31 = f21.split(".")
            for j in range(len(f31)):

                string12 = re.search("|".join(["Sex ", "sex:", "Sex:", "gender", "Patient"]), f31[j])
                if string12 != None and j + 1 < len(f31):
                    s2 = re.findall('Male|Female|MALE|FEMALE|F|M|male|female', f31[j])
                    for i in s2:
                        gender = i

            f53.close()
            # ------------------------------------------------------------------------------

            # ----------------add social and family history to table------------------------

            str27 = folder + "/social_hist.txt"
            f56 = open(str27, "r")
            text = f56.read()
            sents_smoke = []
            sents_alco = []
            # ------------------------------------------------------------------------------

            # ------------------------for smoking-------------------------------------------
            for sent in sent_tokenize(text):
                flag_s = 0
                flag_a = 0
                sent = sent.lower()
                smoke_attr = ["smok", "cigarett"]
                for attr in smoke_attr:
                    if attr in sent:
                        flag_s = 1
                        if flag_s == 1:
                            sents_smoke.append(sent)

                if 'alcohol' in sent:
                    flag_a = 1
                    if flag_a == 1:
                        sents_alco.append(sent)

                list_words1 = ['no', 'never', 'deni']
                list_freq = ['ppd', 'pack/day']
                count1_s = 0
                count2_s = 0
                smoke_var = "NA"
                alcohol_var = "NA"
                freq1 = []
                for sent in sents_smoke:
                    for f in list_freq:
                        freq = re.findall(r'[0-9]+\s*' + f, sent, re.I | re.M)
                        freq1.append(freq)
                    flag_s = 1
                    for word in list_words1:
                        if word in sent:
                            flag_s = 0
                    if flag_s == 0:
                        count1_s += 1
                    else:
                        count2_s += 1

            if count1_s > count2_s:
                smoke_var = "no"
            if count2_s > count1_s:
                smoke_var = "yes"
            f = '-'
            for i in range(0, len(freq1)):
                if freq1[i] != []:
                    f = freq1[i]

            # ------------------------------------------------------------------------------

            # -------------------------------for alcohol------------------------------------
            count1_a = 0
            count2_a = 0
            for sent in sents_alco:
                flag_a = 1
                for word in list_words1:
                    if word in sent:
                        flag_a = 0
                if flag_a == 0:
                    count1_a += 1
                else:
                    count2_a += 1
            if count1_a > count2_a:
                alcohol_var = "no"
            if count2_a > count1_a:
                alcohol_var = "yes"
            f56.close()
            # ------------------------------------------------------------------------------

            # ---------------------------family history-------------------------------------
            str28 = folder + "/fam_hist.txt"
            f57 = open(str28, "r")
            text = f57.read()
            diseases = WordListCorpusReader('E:\BE\RetrievalCorpora', ['diseases.txt'])
            list_diseases = diseases.words()
            members = ["mother", "father"]
            sents_fam = []
            for sent in sent_tokenize(text):
                sent = sent.lower()
                for member in members:
                    if member in sent:
                        sents_fam.append(sent)
            # print(sents_fam)
            disease_m = "NA"
            disease_f = "NA"
            for sent in sents_fam:
                for member in members:
                    if member in sent and member == "mother":
                        for disease in list_diseases:
                            if disease in sent:
                                disease_m = disease

                    if member in sent and member == "father":
                        for disease in list_diseases:
                            if disease in sent:
                                disease_f = disease
            cur.execute(
                """insert into client_records values(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) ON CONFLICT DO NOTHING;""",
                (policy_num, dob, gender, visit, hei1, wei1, bp, smoke_var, f[0], alcohol_var, disease_m, disease_f,))
            conn.commit()
            f57.close()
            # ---------------------------------------------------------------------------

            # --------------split contents of text file into separate lists---------------
            str29 = folder + "/final_vector.txt"
            f58 = open(str29, "r")
            text = f58.read()
            list1 = text.splitlines()
            col_disease = []
            col_var = []
            j = 0
            for i in range(0, len(list1)):
                x = list1[i].split(':')
                y = x[0]
                z = x[1]
                col_disease.append(y)
                col_var.append(z)

            for i in range(0, len(col_disease)):
                d = col_disease[i]
                v = col_var[i]
                insert_statement = """update client_records set "%s" = ('%s') where policy_num = ('%s');""" % (
                d, v, policy_num)
                cur.execute(insert_statement)
            conn.commit()
            f58.close()

        #----------------------------------------------------------------------

        # connection = MongoClient()
        # # connect to the database newdb3
        # db = connection["newdb3"]
        #
        # dirs = os.listdir("E:\BEProject")
        # print(dirs)
        # print(dirs[0])
        # print(os.path.isdir(dirs[0]))
        # for i in range(len(dirs)):
        #     if os.path.isdir(dirs[i]) == True:
        #         folder = dirs[i]
        #         coll = db[folder]  # create new collection for each folder
        #         str18 = "E:\BE_Project" + "/" + folder
        #         dirs1 = os.listdir(str18)
        #         for j in range(len(dirs1)):
        #             if os.path.isdir(dirs1[j]) == False:
        #                 filename = dirs1[j]
        #                 str19 = str18 + "/" + filename
        #                 content = open(str19, 'rb').read().decode('utf8', 'ignore')
        #                 record1 = {'filename': filename,
        #                            'content': content}  # insert all files from folder as key value pairs
        #                 coll.insert(record1)

